"""
build_consolidated_report.py

Builds the single consolidated Word report pulling together:
  - findings_report.tex           (Report 1: full-capacity, 24,744-param, T=50/20ep)
  - findings_report_small.tex     (Reduced-capacity, both regimes T50/10ep + T30/5ep)
  - findings_report_t30.tex       (Reduced-capacity, T30/5ep, more careful seed-1 revision)
  - findings_report_experiments234.tex (Priority Programme Experiments 2-4)
  - Fix 1 (D_l,t layer-relative epsilon) and Fix 2 (E90 multi-draw averaging)

into one docx, per the 12-part structure (Parts 0-11) requested.

No new experiments are run here -- this script only assembles existing text,
tables, and PNG figures already on disk into a single document.
"""
import os
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.expanduser("~/snn")
MAX_WIDTH_IN = 6.4  # usable page width at 1" margins on letter/A4-ish

doc = Document()

# ---------------------------------------------------------------------------
# Style setup
# ---------------------------------------------------------------------------
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

for i, size in zip(range(1, 4), [20, 15, 12.5]):
    h = doc.styles[f"Heading {i}"]
    h.font.size = Pt(size)
    h.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    h.font.bold = True

sections = doc.sections
for s in sections:
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)
    s.top_margin = Inches(0.9)
    s.bottom_margin = Inches(0.9)


def add_page_break():
    doc.add_page_break()


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


def h3(text):
    doc.add_heading(text, level=3)


def p(text="", italic=False, bold=False, size=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.italic = italic
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return para


def bullet(text, bold_lead=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_lead:
        r = para.add_run(bold_lead)
        r.bold = True
        para.add_run(text)
    else:
        para.add_run(text)
    return para


def numbered(text, bold_lead=None):
    para = doc.add_paragraph(style="List Number")
    if bold_lead:
        r = para.add_run(bold_lead)
        r.bold = True
        para.add_run(text)
    else:
        para.add_run(text)
    return para


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def keypoint(text):
    """A shaded single-cell 'table' used as a callout box, mirroring the
    tcolorbox 'keypoint' boxes in the original LaTeX reports."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    shade_cell(cell, "F2F2F2")
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.italic = True
    tbl.autofit = True
    return tbl


def make_table(headers, rows, col_widths=None, small=False, bold_rows=None):
    bold_rows = bold_rows or set()
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        if small:
            run.font.size = Pt(9)
    for ridx, row in enumerate(rows):
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            if small:
                run.font.size = Pt(9)
            if ridx in bold_rows:
                run.bold = True
    return tbl


def figure(relpath, caption, width_in=None, exists_note=None):
    """Embed a PNG at relpath (relative to ~/snn) with aspect-ratio-correct
    sizing, capped at MAX_WIDTH_IN, plus a caption line naming the file."""
    full = os.path.join(BASE, relpath)
    if not os.path.exists(full):
        warn = doc.add_paragraph()
        r = warn.add_run(f"[FIGURE NOT ON DISK: {relpath} -- {exists_note or 'not generated in the source report; noted here rather than silently skipped.'}]")
        r.italic = True
        r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
        return
    im = Image.open(full)
    w, h = im.size
    target_w = width_in or MAX_WIDTH_IN
    target_w = min(target_w, MAX_WIDTH_IN)
    doc.add_picture(full, width=Inches(target_w))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"{caption}  [{relpath}]")
    r.italic = True
    r.font.size = Pt(9)


def add_toc():
    para = doc.add_paragraph()
    run = para.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-2" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "Right-click and choose 'Update Field' to build the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_sep)
    r_element.append(fld_text)
    r_element.append(fld_end)


print("Setup complete, building document...")

# ===========================================================================
# TITLE PAGE
# ===========================================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Rotation Invariance in Spiking and Analog\np4-Equivariant CNNs")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Consolidated Findings Report")
r.font.size = Pt(15)
r.italic = True

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Full-Capacity (24,744-Param) and Reduced-Capacity (2,420-Param) Studies,\n"
                 "the Priority Experimental Programme (Experiments 2-4), and Two Post-Hoc Corrections")
r.font.size = Pt(11)
r.italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Sahil Singh\nCMI Summer Research Internship\nSupervisor: Prof. K.V. Subrahmanyam\n\n"
             "Consolidated: August 2026").font.size = Pt(11)

add_page_break()

h1("Table of Contents")
add_toc()
add_page_break()

# ===========================================================================
# PART 0: HOW TO READ THIS REPORT
# ===========================================================================
h1("Part 0 -- How to Read This Report")

p("This document consolidates four previously separate reports and two post-hoc corrections into a "
  "single, internally consistent record. It reproduces no new experiments and trains no new models -- "
  "every table and figure below reads from results already on disk in ~/snn/ at the time of writing.")

p("The four source reports being merged:")
bullet("Report 1 (“full-capacity companion report”): 7-layer, 10-channel, 24,744-parameter "
       "SpikingP4CNNSmall, T=50 simulation timesteps, 20 training epochs, 4 spike-placement "
       "configurations, 3 seeds. File: findings_report.tex / findings_report.pdf.",
       bold_lead="Report 1. ")
bullet("Report 2 (“reduced-capacity companion report”): 5-layer, 4-channel, 2,420-parameter "
       "SpikingP4CNNSmall, run under two independent regimes (T=50/10ep and T=30/5ep), 5 spike-placement "
       "configurations, 3 seeds. Files: findings_report_small.tex (both regimes) and "
       "findings_report_t30.tex (T=30 regime, revised after seed 1 arrived, more cautious framing of "
       "the group-angle-error finding).",
       bold_lead="Report 2. ")
bullet("Report 3 (“Priority Experimental Programme, Experiments 2-4”): Experiments 2 "
       "(spike placement vs. accuracy/E90), 3 (layer/time representation distance D_{ℓ,t}), and 4 "
       "(Fourier decomposition of the rotation response), run on the same reduced-capacity checkpoints "
       "as Report 2, both regimes, no retraining. File: findings_report_experiments234.tex.",
       bold_lead="Report 3. ")
bullet("A code-vs-methodology audit of three specific anomalies raised across Reports 2 and 3, followed "
       "by two concrete fixes applied to the analysis code: Fix 1 (a layer-relative epsilon and "
       "floor-masking for the D_{ℓ,t} metric) and Fix 2 (multi-draw Monte-Carlo averaging for the "
       "E90 direct-equivariance-error statistic). No model, training, or architecture code was touched "
       "by either fix -- both are corrections to analysis/metric code only.",
       bold_lead="The corrections. ")

keypoint("Superseding notice: the E90 numbers in Part 5 and the D_{ℓ,t} numbers in Part 6 supersede "
         "the equivalent single-draw / fixed-epsilon numbers reported in Reports 2 and 3. Every place "
         "this applies is flagged inline with the phrase “supersedes” rather than left implicit. "
         "The mechanism behind both corrections is explained once, in Part 1's closing paragraph, and "
         "referenced from Parts 5/6 as “the correction described in Part 1” rather than "
         "re-explained.")

p("Reading order: Parts 1-2 establish shared setup and the accuracy backdrop against which every other "
  "finding should be read. Parts 3-8 present the substantive findings, one experiment family per part, "
  "each explicitly stating which report(s) and which regime(s) it draws from. Part 9 is a permanent "
  "settled-questions record of the code-vs-methodology audit, so this project's anomalies are not "
  "re-investigated by a future reader. Part 10 synthesizes at the confidence level each finding actually "
  "supports. Part 11 consolidates limitations and next steps against the current deadlines.")

add_page_break()

# ===========================================================================
# PART 1: SETUP
# ===========================================================================
h1("Part 1 -- Setup")

h2("1.1 Architecture")

p("Both variants share the same qualitative design: a stack of p4-equivariant convolutions "
  "(P4ConvZ2 lifting the input into the p4 group, then P4ConvP4 layers operating within it), "
  "P4BatchNorm2d (sharing statistics and affine parameters across the 4 orientation channels -- see the "
  "callout below), a final group-pooling step (max over the 4 orientation channels) producing an "
  "orientation-invariant feature map, and a linear classifier. Full implementation: gconv_small.py "
  "(P4ConvZ2 / P4ConvP4 / P4BatchNorm2d / P4GroupPool) and p4cnn_small.py / spiking_p4cnn_small.py "
  "(the analog and spiking network definitions).")

make_table(
    ["", "Full-capacity variant", "Reduced-capacity variant"],
    [
        ["Layers", "7", "5"],
        ["Channels per layer", "10", "4"],
        ["Parameters", "24,744", "2,420 (10.2x fewer)"],
        ["Kernels", "six 3x3 + one 4x4 (valid conv)", "3x3 throughout, two 2x2 maxpools"],
        ["Used in", "Report 1 only", "Reports 2 and 3"],
    ],
)

keypoint("P4BatchNorm2d exists because standard BatchNorm assigns each of the 4 orientation channels "
         "independent running statistics, which breaks the architecture's equivariance guarantee once "
         "trained upright-only (verified: untrained equivariance error ~1e-6; after training with "
         "standard BatchNorm, ~0.1-0.2). Sharing statistics and affine parameters across the 4 "
         "orientation channels restores equivariance to ~1e-6 even after training. This fix is used "
         "throughout every result in this document -- it is a settled, resolved issue, listed here for "
         "completeness and cross-referenced from Part 9.")

h2("1.2 Task")
p("4-class rotated-MNIST: digits {3, 4, 8, 9}, chosen because none is rotationally self-symmetric under "
  "90 degrees/180 degrees. Models are trained on upright (0 degree) images only and evaluated on a "
  "24-angle sweep, theta in {0, 15, 30, ..., 345} degrees, on the held-out test split (Reports 1-2) or "
  "a fixed probe set (Report 3 -- see Part 5/6/7 setup notes).")

h2("1.3 Spike-Placement Configurations and Which Report Trained Which")

make_table(
    ["Configuration", "Spiking layers", "Full-capacity (Report 1)", "Reduced-capacity (Reports 2-3)"],
    [
        ["Analog / None", "none", "yes (“Analog”)", "yes (“None”)"],
        ["Alternating", "interleaved through depth (2,4,6 full-cap; 2,4 reduced)", "yes", "yes"],
        ["Full Spiking", "all layers", "yes", "yes"],
        ["Hybrid Late", "final layer only", "yes", "yes"],
        ["Hybrid Early", "first 1-2 layers", "not trained -- see note", "yes"],
        ["Hybrid Mid", "middle layers", "not trained -- see note", "does not exist at depth 5"],
    ],
)
p("Note: Report 1 explicitly states Hybrid Early and Hybrid Mid were implemented in the codebase but "
  "not yet trained under its corrected setup, so Report 1 says nothing about them. Both were trained "
  "for the reduced-capacity model (Reports 2-3); Hybrid Mid has no reduced-capacity analogue since "
  "a 5-layer network has no distinct “middle” block the way a 7-layer network does.",
  italic=True)

h2("1.4 Training Regimes and Seed Counts")

make_table(
    ["Regime", "Model", "T (timesteps)", "Epochs", "Params", "Seeds", "Used by"],
    [
        ["Full-capacity", "SpikingP4CNNSmall (7L/10ch)", "50", "20", "24,744", "3 (0,1,2), full repr. analysis", "Report 1"],
        ["Regime A", "SpikingP4CNNSmall (5L/4ch)", "50", "10", "2,420", "3 (0,1,2)", "Reports 2, 3"],
        ["Regime B", "SpikingP4CNNSmall (5L/4ch)", "30", "5", "2,420", "3 trained; repr. analysis on seeds 0-1 only (seed 2 intentionally incomplete, by explicit instruction)", "Reports 2, 3"],
    ],
)
p("All regimes: Poisson/Bernoulli rate coding of the input, Adam, learning rate 1e-3, batch size 64, "
  "best checkpoint by 0-degree test accuracy, gpuserver.cmi.ac.in.", italic=True)

h2("1.5 The Two Corrections (Referenced Throughout as “the correction described in Part 1”)")

p("Both corrections were produced by a code-vs-methodology audit that investigated three specific "
  "anomalies raised in Reports 2-3 (full detail and verdicts in Part 9). Two of the three anomalies "
  "turned out to be genuine methodological artifacts worth correcting in the analysis code -- not bugs "
  "in the model, training, or the p4-equivariant conv/BatchNorm implementation, which were independently "
  "verified to have no mode-dependent branching.")

numbered("A single fixed epsilon (1e-8) is added to a reference-norm denominator when computing the "
         "layer/time representation distance D_{ℓ,t}(θ) (Report 3, Experiment 3). At the "
         "network's narrowest layer (layer 5, 16 units), the true reference norm is legitimately exactly "
         "zero for a nontrivial fraction of (timestep, image) cells -- dividing by the 1e-8 floor rather "
         "than a real signal produced the reported ~1e7 blowups. Fix: a per-layer epsilon "
         "(max(1e-8, 1% of the median nonzero reference norm at that layer), computed once from the "
         "reference batch) plus explicit exclusion of any cell whose raw reference norm falls below "
         "1e-4 from the reported mean. Applied in representations_small.py's layer_time_distance() "
         "and rerun via the new plot_experiment3_layertime_fixed.py.",
         bold_lead="Fix 1 (D_{ℓ,t} epsilon). ")
numbered("The direct equivariance-error computation (E90, used for Experiment 2's “does spike "
         "placement matter” ordering claim) resets the encoding RNG to the same seed value for both "
         "the upright image x and its rotation R_theta(x) before drawing the Poisson/Bernoulli spike "
         "encoding. Because rotation moves image content between tensor positions, sharing a seed value "
         "reproduces the same raw random stream but does not reproduce a matched spike encoding under "
         "rotation -- it only gives run-to-run reproducibility, not the variance cancellation the "
         "convention seems to have intended. A single draw's E90 is therefore dominated by Monte-Carlo "
         "spike-sampling variance, not a stable property of the trained model (confirmed empirically: "
         "re-drawing the encoding seed alone, with no rotation or config change, reproduces the reported "
         "group-angle magnitudes and ordering). Fix: draw the encoding 8 independent times per "
         "(model seed, config), each with a genuinely different encoding seed "
         "(seed*1000 + draw_index, never reused), and report the mean and std across draws instead of "
         "a single value. Added as equivariance_error_multidraw() in representations_small.py and run "
         "via the new plot_experiment2_summary_multidraw.py.",
         bold_lead="Fix 2 (E90 multi-draw averaging). ")

p("Both fixes touch analysis/metric code only (representations_small.py plus two new driver scripts); "
  "gconv_small.py, p4cnn_small.py, spiking_p4cnn_small.py, and every training script are untouched, and "
  "no model was retrained.", italic=True)

add_page_break()

# ===========================================================================
# PART 2: BASELINE ACCURACY
# ===========================================================================
h1("Part 2 -- Baseline Accuracy (All Configs, All Regimes)")

p("Across every regime in this project, the qualitative story is the same: once training is done "
  "properly (P4BatchNorm2d, adequate epochs), rotation-accuracy curves for every spike-placement "
  "configuration sit within a few percentage points of each other, with substantial seed-to-seed "
  "error-bar overlap at the hardest angles. Accuracy alone does not distinguish spike placement in any "
  "regime tested. The numeric tables below make this precise per regime.")

h2("2.1 Full-Capacity (Report 1): T=50, 20 Epochs, 24,744 Params")
p("Report 1 does not publish a per-configuration numeric accuracy table -- only the comparison figure "
  "below plus qualitative text are available in that report. Its text states the four configurations' "
  "curves sit within a few points of each other at every angle, with substantial error-bar overlap at "
  "the hardest angles (45, 135, 225, 315 degrees), confirming that two dramatic accuracy-level findings "
  "from an earlier, compute-constrained (T=30, 5 epoch, single-seed, CPU-only) draft of the same project "
  "did not replicate under this corrected setup (Full Spiking's earlier flat 51-57% and Hybrid Late's "
  "earlier collapse to chance at non-group angles). Under the corrected setup, Full Spiking reaches "
  "~99% upright and ~94% mean accuracy across all angles; Hybrid Late's worst case across all seeds and "
  "angles is 80.9%.", italic=True)
figure("plots/n4_comparison_final.png",
       "Full-capacity: 4-class rotation-angle sweep, mean +/- std over 3 seeds, "
       "Analog / Alternating / Full Spiking / Hybrid Late (Hybrid Early/Mid: not trained, see Part 1.3).")

h2("2.2 Reduced-Capacity, Regime A: T=50, 10 Epochs, 2,420 Params")
make_table(
    ["Configuration", "Best 0-deg acc. (%), mean+-std, 3 seeds", "Mean acc. over 24 angles (%)"],
    [
        ["None", "97.67 +/- 0.48", "84.77"],
        ["Alternating", "97.39 +/- 0.73", "85.62"],
        ["Full Spiking", "96.39 +/- 0.36", "86.30"],
        ["Hybrid Early", "95.14 +/- 3.53", "81.32"],
        ["Hybrid Late", "97.60 +/- 0.24", "85.84"],
    ],
)
figure("plots_small/small_n5_comparison.png",
       "Regime A (T=50, 10 epochs), all 5 configurations, mean +/- std over 3 seeds.")

h2("2.3 Reduced-Capacity, Regime B: T=30, 5 Epochs, 2,420 Params")
make_table(
    ["Configuration", "Best 0-deg acc. (%), mean+-std, 3 seeds", "Mean acc. over 24 angles (%)"],
    [
        ["None", "97.25 +/- 0.34", "84.52"],
        ["Alternating", "96.71 +/- 0.12", "86.87"],
        ["Full Spiking", "95.14 +/- 0.53", "84.31"],
        ["Hybrid Early", "91.39 +/- 5.71", "77.86"],
        ["Hybrid Late", "97.09 +/- 0.08", "82.03"],
    ],
)
figure("plots_small_t30/small_n5_comparison.png",
       "Regime B (T=30, 5 epochs), all 5 configurations, mean +/- std over 3 seeds.")

p("One configuration breaks the “accuracy is uninformative” pattern in both reduced-capacity regimes: "
  "Hybrid Early shows a conspicuously large seed-to-seed standard deviation (+/-3.53 at Regime A, "
  "+/-5.71 at Regime B), driven by a single unstable seed (seed 2: 90.20% and 83.33% best-upright-accuracy "
  "respectively, versus >=95% for seeds 0 and 1 in both regimes). This instability has no full-capacity "
  "analogue since Report 1 never trained Hybrid Early; it is plausibly a capacity effect (at only 4 "
  "channels, Hybrid Early's spiking layers 1-2 sit immediately after the lifting convolution, where the "
  "network has the least redundancy to absorb a bad initialization's interaction with the "
  "surrogate-gradient spiking nonlinearity) and was not investigated further (Part 11).")

add_page_break()

# ===========================================================================
# PART 3: LAYER-WISE REPRESENTATION SIMILARITY (EXPERIMENT 5.1)
# ===========================================================================
h1("Part 3 -- Layer-wise Representation Similarity (Experiment 5.1)")

p("This finding was consistent everywhere it was tested -- full capacity, both reduced-capacity "
  "regimes, every spike-placement configuration including the zero-spiking control -- and is stated "
  "here plainly, without hedging, because the evidence does not call for any.")

keypoint("Representational similarity to the upright image (measured by cosine similarity, CKA, and "
         "SVCCA between h_ell(x) and h_ell(R_theta x)) stays high through the early and middle layers "
         "and collapses only at the final layer(s), in every configuration tested at every capacity, "
         "including Analog / None, which has zero spiking neurons anywhere in the network. Whatever "
         "causes deep-layer representations to become rotation-sensitive is a property of the "
         "p4-equivariant architecture itself (most plausibly the group-pooling and final linear-classifier "
         "stages, the only components not built from strictly equivariant convolutions) -- not something "
         "introduced by spiking dynamics.")

h2("3.1 Full Capacity (Report 1)")
p("CKA similarity to the upright representation stays at or above roughly 0.97-0.99 through layers 1-5 "
  "in every configuration, then drops sharply in layers 6-7 to roughly 0.70-0.85 at the hardest angles "
  "(typically 15, 30, 45, 60 degrees etc., the non-group angles nearest a group angle). Spike placement "
  "modulates the shape somewhat -- Alternating and Full Spiking show a slightly earlier onset of "
  "divergence (visible already around layer 5) relative to Analog and Hybrid Late -- but not the "
  "existence or final location of the collapse: all four configurations land in the same 0.70-0.85 CKA "
  "range by layer 7 at the hardest angles. Cosine similarity is much lower and noisier throughout "
  "(typically 0.3-0.6 even at shallow layers) since it demands exact neuron-by-neuron agreement with no "
  "tolerance for internal reshuffling of which channel encodes which feature; CKA and SVCCA both stay "
  "near 1.0 through layer 5 and diverge only at layers 6-7.")
figure("plots/analog_layerwise_cka.png", "Full capacity, Analog -- CKA similarity to upright vs. layer, one line per test angle.")
figure("plots/alternating_layerwise_cka.png", "Full capacity, Alternating -- same quantity.")
figure("plots/full_layerwise_cka.png", "Full capacity, Full Spiking -- same quantity.")
p("Hybrid Late's standalone aggregate CKA plot (hybrid_late_layerwise_cka.png) was not generated in "
  "Report 1's pass -- explicitly noted there, not silently omitted here. Its per-seed 5.1 figures were "
  "used instead in that report and show the same collapse pattern.", italic=True)

h2("3.2 Reduced Capacity, Both Regimes (Reports 2-3)")
p("With only 5 layers instead of 7 there is less “room” for a late-layer-only collapse to look visually "
  "distinct from a whole-network effect, yet the qualitative pattern reproduces cleanly: similarity to "
  "the upright representation stays high through the early-to-middle layers and drops at the final "
  "layer(s), in every configuration and both training regimes, including None. This is the single most "
  "important replication in Report 2: the architectural (not spiking) origin of the collapse holds at "
  "10.2x less capacity, 40% less simulation time (Regime B), and a quarter of the training epochs "
  "(Regime B).")
figure("plots_small/repr_seed0/similarity_comparison_cka.png",
       "Regime A, seed 0 -- CKA similarity to upright vs. layer, all 5 configurations overlaid.")
figure("plots_small_t30/repr_seed0/similarity_comparison_cka.png",
       "Regime B, seed 0 -- same comparison.")
p("Spike placement's secondary modulation of exactly where the divergence first becomes visible "
  "(Alternating/Full Spiking diverging one layer earlier than None/Hybrid Late, per the full-capacity "
  "finding) is harder to isolate cleanly with only 5 layers to distribute the effect across -- this is a "
  "genuine capacity-driven limitation on how finely the secondary claim can be tested at this depth, not "
  "a contradiction of it.", italic=True)

# ===========================================================================
# PART 4: ROTATION SENSITIVITY BY LAYER AND BY TIME (EXPERIMENTS 5.3, 5.4)
# ===========================================================================
add_page_break()
h1("Part 4 -- Rotation Sensitivity, by Layer and by Time (Experiments 5.3, 5.4)")

h2("4.1 5.3: Distribution of R_i Across Layers")
p("R_i = Var_theta(a_i(R_theta x)), the rotation sensitivity of neuron i over the 24 test angles. In "
  "every configuration and every capacity, the final layer sits at or near the top of the R_i ranking, "
  "consistent with Part 3's identification of the last layer(s) as where invariance breaks down. The "
  "shape of the climb to the final layer differs meaningfully by spike placement, at full capacity:")
make_table(
    ["Configuration", "Shape of R_i climb to final layer (full capacity)"],
    [
        ["Alternating", "sharp, order-of-magnitude jump only at the final layer: median R_i ~10-100 "
                         "for the earlier layers, jumping to ~400-800 at layer 7"],
        ["Full Spiking", "gradual, full-depth climb: median R_i rises roughly monotonically from ~10 "
                          "(layer 1) to ~90-100 (layer 7), no single sharp jump"],
        ["Hybrid Late", "nearly flat across all layers: median R_i in a comparatively narrow band "
                         "(~10-80) at every depth, including the final layer"],
        ["Analog", "gradual full-depth climb similar in shape to Full Spiking, but ~1000x smaller in "
                    "absolute scale (~0.01-0.05 throughout)"],
    ],
)
figure("plots/full_rotation_sensitivity.png", "Full capacity, Full Spiking -- R_i distribution by layer (log scale).")
p("Hybrid Late's standalone aggregate boxplot (hybrid_late_rotation_sensitivity.png) was likewise not "
  "generated in Report 1's pass; its per-seed 5.3 figures were used instead there.", italic=True)

p("Reduced capacity: the same qualitative fingerprint (sharp late jump for Alternating, gradual climb "
  "for Full Spiking/None, flatter profile for Hybrid Late) reproduces, but is harder to resolve cleanly "
  "with only 5 layers to distribute the pattern across. This is consistent with, not contradicted by, "
  "the full-capacity finding -- the reduced-capacity report is underpowered to cleanly distinguish the "
  "shapes at this depth, not evidence the shapes actually converge.", italic=True)
figure("plots_small/repr_seed0/rotsens_median_comparison.png",
       "Regime A, seed 0 -- median R_i per layer, all 5 configurations overlaid (log scale).")

keypoint("R_i's absolute scale depends on whether the underlying activation is a continuous ReLU output "
         "or a spike count, so the fair cross-configuration comparison is the shape of the R_i-vs-layer "
         "curve, not the raw magnitude. Within a single configuration, comparisons across layers or "
         "seeds are on safe footing.")

h2("4.2 5.4: Rotation Sensitivity Through Time")
p("R_i(t): the same quantity computed from the instantaneous per-timestep spike output alone, plotted "
  "against timestep t. In every configuration, every seed, and every capacity/regime tested, R_i(t) "
  "rises quickly over the first few timesteps and then plateaus for the remainder of the simulation, "
  "with no visible downward trend. This directly contradicts the temporal-contraction hypothesis that "
  "motivated this experiment (that temporal averaging should progressively reduce sensitivity, "
  "R_i(t) decaying toward 0 as spikes accumulate) -- sensitivity is set almost immediately and then "
  "stays fixed.")
figure("plots/full_seed0/full_seed0_5.4_rotation_sensitivity_through_time.png",
       "Full capacity, Full Spiking, seed 0 -- R_i(t) vs. timestep, all 7 layers.")
figure("plots_small/repr_seed0/rotsens_temporal_full.png",
       "Regime A, Full Spiking, seed 0 -- median R_i(t) vs. timestep, all 5 layers.")
figure("plots_small_t30/repr_seed0/rotsens_temporal_full.png",
       "Regime B, Full Spiking, seed 0 -- same quantity, T=30.")
p("Layer 7 (full capacity) plateaus visibly higher than other layers in Alternating and Full Spiking, "
  "consistent with Part 3; in Hybrid Late the plateau levels across layers are closer together, "
  "consistent with that configuration's flatter layer-wise sensitivity profile. Analog has no "
  "meaningful 5.4 plot at full capacity (T=1, no timestep axis). The plateau-not-decay result "
  "reproduces exactly at reduced capacity, in both regimes, with no exceptions found in any "
  "configuration or seed examined.", italic=True)

# ===========================================================================
# PART 5: DIRECT EQUIVARIANCE ERROR (EXPERIMENT 5.8 / EXPERIMENT 2's E90)
# ===========================================================================
add_page_break()
h1("Part 5 -- Direct Equivariance Error (Experiment 5.8 / Experiment 2's E90)")

p("This is the part of the project where the most care is required, because two genuinely different "
  "quantities were measured under the same name (“equivariance error”) across reports, and one of them "
  "(the reduced-capacity E90 value) needed a genuine correction. This section is deliberately structured "
  "in four steps so none of that gets blurred together.")

h2("5.1 (a) Two Different Quantities, Not Directly Comparable")
p("Both quantities are legitimate, correctly-implemented computations of the relative error "
  "||f(R_theta x) - f(x)|| / ||f(x)||, which the p4 architecture guarantees is exactly 0 at the four "
  "group angles {0, 90, 180, 270} by construction. They differ in two ways that turn out to matter a "
  "great deal:")
make_table(
    ["", "Report 1 (full capacity)", "Reports 2-3 (reduced capacity, E90)"],
    [
        ["Space", "raw accumulated-spike-count logits (no softmax)", "softmax class-probability vector"],
        ["Angle reduction", "peak (max) over all 24 test angles", "value at exactly theta=90 degrees"],
    ],
)
p("The audit confirmed directly (not just by formula inspection) that this difference, not model "
  "capacity, drives the apparently contradictory orderings below: recomputing BOTH formulas on the SAME "
  "reduced-capacity checkpoints reproduces Report 1's ordering under the raw-logit/peak-over-angles "
  "formula (None highest at 0.417, Full Spiking lowest at 0.270) and reproduces the opposite ordering "
  "under the softmax/value-at-90 formula (Full Spiking highest at 0.498) -- on identical models. "
  "Softmax compresses outputs onto a bounded probability simplex, which changes both which angle "
  "maximizes the disagreement and how the denominator scales with each configuration's confidence/"
  "calibration, which differs systematically by configuration. Sections (b) and (c) below report each "
  "quantity as its own report found it; they are not reconciled into one ordering (see (d)).",
  italic=True)

h2("5.2 (b) Full-Capacity Finding: Peak Equivariance Error (Report 1, As-Is)")
p("Quantity: raw-logit relative error, peak over all 24 angles. Every configuration's error touches (or "
  "comes extremely close to) 0 at exactly the four group angles and rises to a local maximum roughly "
  "halfway between them, in every configuration and every seed -- the expected sawtooth, and a strong "
  "sanity check that P4BatchNorm2d's equivariance fix functions correctly in fully trained models.")
figure("plots/analog_seed0/analog_seed0_5.8_equivariance_error.png",
       "Full capacity, Analog, seed 0 -- equivariance error vs. angle.")
figure("plots/full_seed0/full_seed0_5.8_equivariance_error.png",
       "Full capacity, Full Spiking, seed 0 -- same quantity.")
make_table(
    ["Configuration", "Spiking layers", "Approx. peak equivariance error (mean over seeds)"],
    [
        ["Analog", "0/7", "~0.58-0.59 (highest)"],
        ["Alternating", "3/7", "~0.50-0.60"],
        ["Hybrid Late", "1/7", "~0.35-0.40"],
        ["Full Spiking", "7/7", "~0.33-0.42 (lowest)"],
    ],
)
p("If spiking dynamics simply degraded invariance in proportion to how many layers spike, peak error "
  "should increase monotonically from Analog through Hybrid Late through Alternating through Full "
  "Spiking. The data does not show this: Analog has the highest peak error and Full Spiking the lowest "
  "-- the opposite of the naive “spiking hurts invariance” expectation. Report 1 reports this as an "
  "open, reproducible observation, not an explained result; no experiment run in that report isolates "
  "why this ordering holds.")

h2("5.3 (c) Reduced-Capacity E90 Finding, CORRECTED (Supersedes Reports 2 and 3)")
p("Quantity: softmax-probability relative error, value at exactly theta=90 degrees (“E90”).",
  italic=True)
p("This supersedes two earlier readings of the same quantity. Report 2's single-seed-0 reading first "
  "suggested Full Spiking was uniquely, severely anomalous at the group angles; Report 2's seed-1 "
  "revision (findings_report_t30.tex) walked that back to “noisy across every configuration, no "
  "configuration reproducibly anomalous, n=2 is not enough to tell.” Report 3, using a probe set held "
  "fixed across seeds (a genuine methodological improvement over Reports 2's per-seed-varying probe "
  "set), found a clean monotonic ordering (None < Alternating < Full Spiking) in both regimes and "
  "described it as materially cleaner than Report 2's reading. The correction described in Part 1 "
  "(Fix 2: 8 independent encoding draws per model seed, each with a genuinely different encoding seed) "
  "was applied on top of Report 3's fixed-probe-set protocol, because the audit traced the underlying "
  "noise source to Monte-Carlo spike-sampling variance in the encoding, not to the probe set.")

make_table(
    ["Regime", "None", "Alternating", "Full Spiking"],
    [
        ["A (T=50)", "0.0038 +/- 0.0044", "0.0102 +/- 0.0038", "0.0249 +/- 0.0151"],
        ["B (T=30)", "0.0146 +/- 0.0087", "0.0248 +/- 0.0147", "0.0455 +/- 0.0168"],
    ],
)
p("(Mean +/- std of E90 over 3 model seeds; each seed's own value is itself the mean of 8 independent "
  "encoding draws, so this combines model-seed variance and encoding-draw variance.)", italic=True)
figure("plots_small/experiment2/exp2_e90_multidraw_comparison.png",
       "E90 before (single-draw) vs. after (8-draw mean) the Fix 2 correction, both regimes, mean +/- std over 3 model seeds.")

p("What survives: the mean ordering None < Alternating < Full Spiking holds in both regimes, "
  "independently -- and it is now corroborated two ways rather than one, since it reproduces at T=50 "
  "and T=30 on entirely different checkpoints and is stable to which encoding draw is used, not just to "
  "which model seed happened to be picked. That is a real strengthening of the claim relative to any "
  "single-draw reading.")
p("What does not survive: the word “clean.” The mean +/- std bands overlap substantially between "
  "adjacent configurations at n=3 model seeds in both regimes (e.g. Regime A: None's upper band 0.0083 "
  "already overlaps Alternating's lower band 0.0064; Alternating/Full Spiking overlap even more "
  "heavily). This is not new noise introduced by the correction -- the old single-draw per-seed numbers "
  "already had the same overlap (e.g. Regime A seed 1: Alternating=0.0213 > Full Spiking=0.0177) -- "
  "multi-draw averaging simply makes the overlap explicit and quantified rather than hidden behind a "
  "single point estimate per seed. Directional and cross-regime-consistent is a defensible claim at "
  "n=3 seeds; clean statistical separation between adjacent configurations is not, and would require "
  "more model seeds, not more encoding draws.", bold=False)

h2("5.4 (d) These Two Findings Are Not Reconciled")
keypoint("Part 5.2 (full-capacity, peak-over-angles, raw logits) and Part 5.3 (reduced-capacity, "
         "value-at-90, softmax, multi-draw-corrected) are different measurements on different models. "
         "No attempt is made here to merge them into one ordering or one causal story, because Part "
         "5.1 already demonstrated that switching the formula alone flips the ordering on identical "
         "checkpoints -- any apparent agreement or disagreement between (b) and (c) as currently "
         "computed would not be informative. This is flagged as an explicit open question for future "
         "work: computing BOTH quantities (raw-logit peak-over-angles, and softmax multi-draw value-at-90) "
         "on the SAME set of checkpoints, ideally the full-capacity model, would be needed before any "
         "claim relating them could be made.")

# ===========================================================================
# PART 6: D_{l,t} (EXPERIMENT 3)
# ===========================================================================
add_page_break()
h1("Part 6 -- Layer/Time Representation Distance D_{ℓ,t} (Experiment 3)")

p("D_{ℓ,t}(theta) = E_x[ ||h_{ℓ,t}(R_theta x) - h_{ℓ,t}(x)|| / (||h_{ℓ,t}(x)|| + eps) ], "
  "the representation distance at a single layer and single timestep, evaluated at theta=45 (no exact "
  "group guarantee) and theta=90 (exact C4 guarantee). This section reports the CORRECTED numbers "
  "(the correction described in Part 1, Fix 1) and states plainly that the original raw values "
  "(~1e5 to ~3.4e7) were floor-division artifacts carrying no signal about representation distance, not "
  "a real result to be interpreted.")

h2("6.1 Layer 5, Mean D_{ℓ,t}(45 deg): Before vs. After")
figure("plots_small/experiment3_fixed/exp3_layer5_before_after_all_combos.png",
       "Layer 5 mean D_{45deg}, all 3 configs x 3 seeds x 2 regimes, before (fixed eps=1e-8, log scale) "
       "vs. after (layer-relative eps + floor masking) the Fix 1 correction.")
make_table(
    ["Regime", "Config", "Seed", "Mean BEFORE", "Mean AFTER", "Cells excluded"],
    [
        ["T50", "None", "0", "5.2e5", "1.03", "1.0%"],
        ["T50", "None", "1", "1.3e6", "1.13", "6.6%"],
        ["T50", "Alternating", "0", "2.1e6", "1.16", "7.6%"],
        ["T50", "Full Spiking", "0", "8.9e6", "1.18", "9.5%"],
        ["T30", "Full Spiking", "0", "9.6e6", "1.12", "12.8%"],
    ],
    small=True,
)
p("Full 18-combination table is in exp3_eps_diagnostics.csv under each regime's experiment3_fixed/ "
  "folder. Every blown-up value collapses to the 0.85-1.4 range once the floor artifact is removed; "
  "seeds that were already sane (e.g. None seed 2: 0.87 before and after) barely move -- confirming the "
  "fix targets exactly the broken cells and does not disturb the ones that were already correct.",
  italic=True)

h2("6.2 Visual Before/After: Heatmaps")
figure("plots_small/experiment3/exp3_heatmap_full_seed0.png",
       "BEFORE (original, plots_small/experiment3/): Full Spiking, Regime A, seed 0 -- "
       "D_{ℓ,t}(45deg) heatmap. The deepest layer's color range is dominated by floor-division spikes.")
figure("plots_small/experiment3_fixed/exp3_heatmap_full_seed0.png",
       "AFTER (corrected, plots_small/experiment3_fixed/): same model/seed/quantity, with the Fix 1 "
       "layer-relative epsilon and floor masking applied.")

h2("6.3 Collapse-Frequency-by-Configuration Finding, Corrected Caveat")
p("Report 3 originally described the layer-5 floor-collapse frequency as “rare in None, frequent in "
  "Alternating, near-constant in Full Spiking,” read from single-seed data. With the corrected metric "
  "computed across all 3 seeds and both regimes, the directional claim survives on average but the "
  "“near-constant” language does not -- it overstates what a single seed showed.")
make_table(
    ["Regime", "None", "Alternating", "Full Spiking"],
    [
        ["T50, mean % cells excluded (theta=45)", "2.5% +/- 3.6", "3.1% +/- 4.0", "3.8% +/- 5.0"],
        ["T30, mean % cells excluded (theta=45)", "2.1% +/- 3.1", "0.7% +/- 1.1", "4.7% +/- 7.0"],
    ],
)
p("On average, across seeds, Full Spiking does have the highest floor-collapse rate in both regimes, "
  "consistent with more spiking layers upstream propagating more sparsity into the narrow final layer. "
  "But this is not a clean per-seed pattern: Full Spiking seed 0 hits 9.5-12.8% exclusion while Full "
  "Spiking seed 1 hits only 0.5-0.6%, comparable to None. The correct framing is “higher on average, "
  "consistent with the number of spiking layers, but with seed-to-seed variance as large as the effect "
  "itself” -- not “near-constant.” This is the same seed-variance caution that recurs elsewhere in this "
  "project (Part 9, Part 10).", italic=True)
figure("plots_small/experiment3_fixed/exp3_lines_full_seed0.png",
       "Regime A, Full Spiking, seed 0 (corrected) -- D_{ℓ,t}(45deg) solid vs. D_{ℓ,t}(90deg) "
       "dashed (layer 5 only), log scale.")

p("The upstream-sparsity-propagation hypothesis first proposed in Report 3 -- that sparsity introduced "
  "by an upstream spiking layer propagates forward and makes a downstream analog layer's pre-activation "
  "more likely to be exactly zero at a given timestep, rather than the layer's own activation type "
  "being what matters -- remains a plausible, mechanistically motivated reading of the corrected data, "
  "but is still not directly verified (Part 11).")

# ===========================================================================
# PART 7: FOURIER ANALYSIS (EXPERIMENT 4)
# ===========================================================================
add_page_break()
h1("Part 7 -- Fourier Analysis of the Rotational Response (Experiment 4)")

p("This experiment needs no correction and is the strongest, most seed-stable result in the project. "
  "It is presented here as reported in Report 3, unchanged, as the headline confirmatory result of the "
  "whole programme.")

p("Setup: for None, Alternating, Full Spiking, all 3 seeds, both regimes, the classification margin "
  "M(theta) = z_y(R_theta x) - max_{c != y} z_c(R_theta x) and accuracy A(theta) are evaluated at "
  "1-degree resolution (360 angles) on a fixed 32-image probe set, using raw accumulated-spike-count "
  "logits (not softmax). M(theta) is Fourier-decomposed up to k_max=20, giving E_noninv (non-constant "
  "energy fraction) and E_forbidden (energy at k not divisible by 4 -- architecturally, this should be "
  "~0 for an exact p4 network).")

figure("plots_small/experiment4/exp4_rotation_response_full.png",
       "Regime A, Full Spiking -- classification margin M(theta), all 3 seeds + mean, 1-degree resolution.")
figure("plots_small/experiment4/exp4_fourier_spectrum_full.png",
       "Regime A, Full Spiking -- Fourier spectrum magnitude vs. k; red = allowed (k a multiple of 4), gray = forbidden.")

p("The spectrum is unambiguous: every forbidden frequency sits at a uniform noise floor "
  "(sqrt(a_k^2+b_k^2) <~ 0.03) indistinguishable from zero, while the allowed frequencies show a clear, "
  "decaying spectrum dominated by k=4 (~2-3), with smaller but clearly nonzero contributions at "
  "k=8,12,16,20. This is the clearest confirmation in the entire project of the p4 architecture's exact "
  "discrete-symmetry guarantee, directly in Fourier space rather than through an accuracy or "
  "representation-similarity proxy.")

make_table(
    ["Regime", "Config", "a0", "E_noninv", "E_forbidden", "|coef k=4|", "|coef k=8|"],
    [
        ["A", "None", "7.11+/-2.62", "0.278+/-0.094", "9.2e-5+/-3.9e-5", "4.07+/-0.74", "0.63+/-0.28"],
        ["A", "Alternating", "7.80+/-1.88", "0.228+/-0.058", "1.11e-4+/-5.5e-5", "4.08+/-0.34", "0.36+/-0.34"],
        ["A", "Full Spiking", "5.94+/-1.61", "0.192+/-0.086", "1.18e-4+/-3.8e-5", "2.70+/-0.51", "0.48+/-0.11"],
        ["B", "None", "5.73+/-0.96", "0.252+/-0.095", "1.56e-4+/-2.6e-5", "3.31+/-1.16", "0.37+/-0.15"],
        ["B", "Alternating", "6.27+/-0.99", "0.174+/-0.064", "1.33e-4+/-4.9e-5", "2.77+/-0.23", "0.22+/-0.14"],
        ["B", "Full Spiking", "4.07+/-0.66", "0.201+/-0.050", "2.84e-4+/-6.7e-5", "2.01+/-0.44", "0.22+/-0.18"],
    ],
    small=True,
)
figure("plots_small/experiment4/exp4_enoninv_eforbidden.png",
       "Regime A -- E_noninv (left) vs. E_forbidden (right), boxplot + individual seeds, all 3 configurations.")

keypoint("E_forbidden is architectural, not learned: all three configurations cluster tightly in the "
         "same narrow band (~6e-5 to ~2.8e-4) across both regimes, with heavily overlapping seed "
         "distributions -- exactly the “essentially independent of the random seed” prediction this "
         "experiment was designed to test. E_noninv is NOT architectural: it varies more across "
         "configurations and seeds (0.17-0.28), with a mild trend toward lower non-invariant energy as "
         "more layers spike, visible for the k=4 coefficient specifically (None > Alternating > Full "
         "Spiking in both regimes) though not fully consistent for E_noninv itself. This is exactly the "
         "architecture-determines-symmetry / training-determines-learned-invariance-strength "
         "distinction the experimental programme set out to isolate.")

# ===========================================================================
# PART 8: COMPUTE-BUDGET FINDINGS (TRUNCATED-T', EXPERIMENT 5.7)
# ===========================================================================
add_page_break()
h1("Part 8 -- Compute-Budget Findings: Truncated-T' Accuracy (Experiment 5.7)")

p("A single model, trained once at its regime's full T, has its recorded spike train truncated to the "
  "first T' timesteps at evaluation time and a classification decision decoded from only those, "
  "without retraining. This tests whether rotation-robust accuracy structure emerges progressively as "
  "spikes accumulate, and how quickly.")

h2("8.1 Full Capacity (Report 1)")
figure("plots/full_seed0/full_seed0_5.7_temporal_evolution.png",
       "Full capacity, Full Spiking, seed 0 -- accuracy vs. angle at T'=1,5,10,20,50.")
figure("plots/hybrid_late_seed0/hybrid_late_seed0_5.7_temporal_evolution.png",
       "Full capacity, Hybrid Late, seed 0 -- same T' grid.")
p("Full Spiking's T'=1 curve is flat at ~25-27% (chance) across every angle -- with all 7 layers "
  "spiking, every layer needs several timesteps to integrate enough input spikes to pass a meaningful "
  "signal onward, and this startup delay compounds across all 7 layers. Real W-shaped structure appears "
  "only by T'=10 and is nearly indistinguishable from T'=50 by T'=20. Alternating's T'=1 curve already "
  "shows a weak but genuine W-shape (25-70%, seed/angle-dependent). Hybrid Late's T'=1 curve already "
  "closely tracks the full T'=50 curve's shape (63-99%) -- with only the final layer spiking, six of "
  "seven layers process instantaneously regardless of T'. This gives a consistent, monotonic ordering "
  "across all 9 seed-runs examined (3 configs x 3 seeds): the more layers that spike, the more "
  "timesteps needed before rotation-robust accuracy structure appears.")

h2("8.2 Reduced Capacity, Both Regimes (Reports 2)")
make_table(
    ["Regime", "Config", "T'=1", "T'=5", "T'=10", "T'=20", "T'=50"],
    [
        ["A (T=50, seed 0)", "Alternating", "51.2%", "84.4%", "86.6%", "87.0%", "89.4%"],
        ["A (T=50, seed 0)", "Full Spiking", "31.2%", "71.0%", "82.6%", "84.7%", "85.9%"],
        ["A (T=50, seed 0)", "Hybrid Late", "58.7%", "87.0%", "88.7%", "89.1%", "89.6%"],
        ["A (T=50, seed 0)", "None", "77.1%", "85.9%", "88.4%", "88.7%", "90.5%"],
        ["B (T=30, mean+-std, seeds 0-1)", "Alternating", "45.5+/-6.6%", "79.6+/-12.6%", "83.9+/-9.7%", "85.1+/-9.0%", "84.9+/-9.0%"],
        ["B (T=30, mean+-std, seeds 0-1)", "Full Spiking", "34.3+/-6.9%", "66.1+/-17.4%", "78.6+/-8.9%", "82.0+/-4.8%", "83.3+/-4.3%"],
        ["B (T=30, mean+-std, seeds 0-1)", "Hybrid Late", "60.3+/-9.1%", "76.0+/-7.3%", "79.1+/-6.0%", "80.3+/-6.7%", "81.5+/-6.3%"],
        ["B (T=30, mean+-std, seeds 0-1)", "None", "73.1+/-6.1%", "80.9+/-4.0%", "83.2+/-3.7%", "84.5+/-3.1%", "84.8+/-3.0%"],
    ],
    small=True,
)
figure("plots_small/repr_seed0/truncated_T_by_angle_full.png",
       "Regime A, Full Spiking, seed 0 -- accuracy vs. angle at T'=1,5,10,20,50.")
figure("plots_small/repr_seed0/truncated_T_by_angle_hybrid_late.png",
       "Regime A, Hybrid Late, seed 0 -- same T' grid.")
p("The same “more spiking layers => longer warm-up” ordering reproduces at reduced capacity, in both "
  "regimes.")

keypoint("Practical implication, consistent across every capacity and regime tested: T'=10 already "
         "recovers most of the T'=50 accuracy everywhere (full capacity: visually indistinguishable by "
         "T'=20; reduced capacity Regime A: >=94% of T'=50 accuracy for every configuration; reduced "
         "capacity Regime B: 93-99% of the T=30-clamped T'=50 accuracy for every configuration). Models "
         "trained and evaluated at the full T throughout this project may be using substantially more "
         "simulation time than necessary at decode time.")

# ===========================================================================
# PART 9: CODE VS METHODOLOGY AUDIT (PERMANENT RECORD)
# ===========================================================================
add_page_break()
h1("Part 9 -- What Was Ruled Out: Code vs. Methodology Audit")

p("This is a permanent settled-questions record of every anomaly investigated during this project's "
  "code-vs-methodology audit, so a future reader does not need to re-investigate them. For each: what "
  "was suspected, what was ruled out, the verdict (code bug vs. methodological/metric artifact), and "
  "the one-line resolution.")

h2("9.1 P4BatchNorm2d Breaking Equivariance After Training (resolved before this audit)")
p("Suspected: standard BatchNorm's independent per-orientation-channel statistics might interact badly "
  "with upright-only training.", italic=True)
p("Verdict: CONFIRMED CODE ISSUE, already fixed. Untrained equivariance error ~1e-6; after training "
  "with standard BatchNorm, ~0.1-0.2. Resolution: P4BatchNorm2d, sharing statistics/affine parameters "
  "across the 4 orientation channels, restores ~1e-6 post-training. Used throughout every result in "
  "this document. This is the one genuine architecture-level bug found and fixed in this project.",
  bold=True)

h2("9.2 Group-Angle Equivariance-Error Noise (Reports 2-3)")
p("Suspected: at exact group angles (90/180/270 degrees), equivariance error should be ~0 by "
  "construction, but was observed 10-100x higher than expected for some configs, inconsistently "
  "across seeds.", italic=True)
bullet("Rotation-implementation noise (grid_sample vs. torch.rot90). Ruled out: max abs pixel "
       "difference ~2e-6 (float32 noise floor), 4-5 orders of magnitude below the observed 0.01-0.06 "
       "errors; feeding the model the bit-exact rotation instead left the error unchanged to 6 decimal "
       "places.", bold_lead="Ruled out (1): ")
bullet("Mode-dependent code branching in gconv_small.py. Ruled out: P4ConvZ2, P4ConvP4, "
       "P4BatchNorm2d, P4GroupPool contain no reference to mode or spiking_layers anywhere -- one code "
       "path for every configuration, confirmed by direct code reading, not inference.",
       bold_lead="Ruled out (2): ")
p("Verdict: METHODOLOGICAL ARTIFACT, not a code bug. The evaluation code resets the encoding RNG to "
  "the same seed value for x and R_theta(x). Because rotation moves image content between tensor "
  "positions, this reproduces the same raw random stream but not a matched spike encoding under "
  "rotation -- it gives run-to-run reproducibility, not the noise-cancellation the convention seems to "
  "have intended. Confirmed empirically: re-drawing the encoding seed alone (same model, same probe "
  "images, no rotation or config change) reproduces the exact reported magnitudes and None < "
  "Alternating < Full Spiking ordering. Also discovered along the way: layer 1's input is stochastically "
  "spike-encoded in EVERY configuration, including None -- “Analog” is not literally analog at the "
  "input, only in its internal layers -- which is why None also showed nonzero, seed-inconsistent noise "
  "at group angles rather than being immune to the effect. Resolution: Fix 2 (Part 1), multi-draw "
  "averaging.", bold=True)

h2("9.3 D_{ℓ,t} Reference-Norm Collapse (Report 3, Experiment 3)")
p("Suspected: the deepest layer's D_{ℓ,t}(45deg) values explode to ~1e7 because the denominator "
  "||h_{ℓ,t}(x)|| + eps is near-zero when that timestep's activation is mostly/all zero.",
  italic=True)
p("Verdict: METHODOLOGICAL/METRIC-DESIGN ARTIFACT, not a model or training bug. eps=1e-8 was a fixed "
  "global constant; measured layer-5 reference norms directly on real checkpoints and confirmed 9.5% of "
  "(timestep, image) cells there have a raw norm of EXACTLY zero (the whole 16-unit layer silent that "
  "timestep), so denom = eps exactly and any nonzero numerator explodes the ratio. Layers 1-4 never hit "
  "an exact zero in the same probe set -- the problem is specific to the narrowest, sparsest layer. "
  "Resolution: Fix 1 (Part 1), layer-relative epsilon plus floor masking.", bold=True)

h2("9.4 Report 1 vs. Report 3 Equivariance-Error Direction Reversal")
p("Suspected: Report 1 found peak equivariance error highest for Analog, lowest for Full Spiking; "
  "Reports 2-3 found error at 90 degrees specifically lowest for Analog/None, highest for Full "
  "Spiking -- apparently contradictory.", italic=True)
p("Verdict: NOT A BUG IN EITHER -- genuinely different quantities being compared (see Part 5.1 for the "
  "full formula comparison). Confirmed directly: recomputing both formulas (raw-logit peak-over-angles; "
  "softmax value-at-90) on the IDENTICAL reduced-capacity checkpoints reproduces each report's ordering "
  "under its own formula -- the formula choice alone flips the ranking, not model capacity. Both "
  "implementations correctly compute what they claim to; the issue is that the two numbers were "
  "juxtaposed across reports as if directly comparable “equivariance error” when they are not. Left "
  "unreconciled by design (Part 5.4).", bold=True)

h2("9.5 Evolution of the “Full Spiking Anomaly” Narrative Across Reports")
p("For completeness: the understanding of 9.2 evolved visibly across this project's own reports, and "
  "this evolution is itself worth recording so it is not mistaken for inconsistency. "
  "findings_report_small.tex (seed-0-only reading) proposed “cascaded LIF-threshold amplification of "
  "spike-sampling stochasticity, scaling with spiking-layer count” as the explanation, based on Full "
  "Spiking looking uniquely and reproducibly elevated across its two available regimes. "
  "findings_report_t30.tex (after seed 1 arrived) walked this back: four of five configurations showed "
  "at least one elevated group-angle value depending on seed, so no single configuration could "
  "responsibly be called uniquely anomalous at n=2. The final audit (Part 9.2, this document) confirms "
  "the original mechanism's core intuition was directionally right (more spiking layers does correlate "
  "with more noise, confirmed cleanly via the multi-draw experiment in Part 5.3) but locates the root "
  "cause more precisely (RNG/pixel-position mismatch at the stochastically-encoded input, present in "
  "every configuration including None, not a LIF-threshold-specific cascade) and treats the effect as "
  "quantifiable Monte-Carlo variance rather than a per-configuration “bug.” Later readings are more "
  "reliable than earlier ones here; this document uses the final audit's framing throughout.")

# ===========================================================================
# PART 10: SYNTHESIS
# ===========================================================================
add_page_break()
h1("Part 10 -- Synthesis")

p("Written at the confidence level each finding actually supports -- Part 5's now-corrected E90 "
  "finding is not upgraded back to “clean ordering” language here just because this is the summary.")

numbered("Accuracy alone is uninformative about spike placement, at every capacity and regime tested. "
         "Once training is done properly, all configurations reach rotation-accuracy curves within a "
         "few points of each other, with one reproducible exception (Hybrid Early's large seed-to-seed "
         "variance at reduced capacity, absent at full capacity only because Report 1 never trained it).")
numbered("The layer-wise representational collapse at the deepest layer(s) is a genuine, "
         "unhedged, universal finding: it appears in every configuration at every capacity tested, "
         "including the zero-spiking control, and is a property of the p4-equivariant architecture's "
         "final stages, not of spiking dynamics.")
numbered("Rotation sensitivity through spiking time rises quickly and then plateaus -- it does not "
         "decay -- in every configuration, every seed, and every regime tested, directly contradicting "
         "the project's original temporal-contraction hypothesis. Per-configuration differences in the "
         "shape of the R_i-vs-layer climb (sharp late jump, gradual climb, or flat) are real at full "
         "capacity and reproduce qualitatively but less cleanly at reduced capacity, where 5 layers give "
         "less room to resolve the shapes distinctly.")
numbered("Truncated-T' accuracy shows a clean, reproducible, mechanistically sensible ordering across "
         "every capacity and regime: more spiking layers means a longer warm-up before rotation-robust "
         "accuracy structure appears, and T'=10 recovers most of the full-T accuracy everywhere -- a "
         "genuine, actionable compute-budget finding.")
numbered("Direct equivariance error tells two different, non-reconciled stories depending on exactly "
         "what is measured. At full capacity, peak error over all angles is non-monotonic in spiking "
         "extent (Analog worst, Full Spiking best) and remains unexplained. At reduced capacity, the "
         "corrected (multi-draw) E90 at exactly 90 degrees shows a directionally consistent, "
         "cross-regime-reproducible ordering (None < Alternating < Full Spiking in the mean) that does "
         "NOT achieve clean separation between adjacent configurations at n=3 model seeds -- this is a "
         "real but modest finding, not the “clean ordering” one of its source reports originally "
         "claimed.")
numbered("Experiment 4's Fourier analysis is the project's strongest and most seed-stable confirmatory "
         "result: forbidden-mode energy is architecturally pinned near zero regardless of configuration "
         "or seed, while allowed-mode amplitude (the learned component) varies substantially -- a clean "
         "empirical separation of what the architecture guarantees from what training actually learns.")
numbered("Two of the three anomalies investigated in this project's code-vs-methodology audit were "
         "genuine methodological artifacts (Monte-Carlo spike-encoding noise; a fixed epsilon "
         "inappropriate for a narrow layer), now corrected in the analysis code; none was a bug in the "
         "p4-equivariant conv/BatchNorm implementation itself, which was independently verified to have "
         "no mode-dependent branching. The one genuine historical code bug in this project "
         "(standard BatchNorm breaking equivariance) was found and fixed before this audit began.")
numbered("A recurring methodological lesson across this entire project, worth stating on its own: "
         "single-seed or single-draw readings repeatedly looked cleaner than they turned out to be once "
         "a second seed, a second draw, or a corrected metric arrived (the Full-Spiking-anomaly "
         "narrative, the D_{ℓ,t} collapse-frequency ordering, and the E90 ordering all followed this "
         "pattern). Every finding in this document that still holds at n=3 seeds has been "
         "cross-checked against this pattern before being stated plainly.")

# ===========================================================================
# PART 11: LIMITATIONS AND RECOMMENDED NEXT STEPS
# ===========================================================================
add_page_break()
h1("Part 11 -- Limitations and Recommended Next Steps")

h2("11.1 Limitations, Consolidated Across All Reports")
bullet("Every seed-averaged number in this document uses n=3 model seeds (n=2 for several "
       "reduced-capacity Regime B representation-level quantities, since seed 2's representation "
       "analysis was intentionally left incomplete by explicit instruction). The Priority Programme "
       "recommends n=10 (Experiment 2) and n=5 (Experiment 3); this project has not reached either.",
       bold_lead="Seed counts below programme recommendations. ")
bullet("SVCCA is not on an absolute 0-1 similarity scale -- independent, genuinely dissimilar data "
       "produces a nonzero baseline of roughly 0.2-0.4 at the batch sizes used here. Read SVCCA values "
       "relatively (larger = more similar), never as an absolute invariance measure.",
       bold_lead="SVCCA absolute-scale caveat. ")
bullet("R_i's absolute magnitude depends on whether the underlying activation is a continuous ReLU "
       "output or a spike count (the spiking configurations' R_i sits roughly 1000x larger than "
       "Analog/None's at comparable layers). Only within-configuration comparisons (across layers or "
       "seeds) are on safe footing; cross-configuration comparisons must use shape, not magnitude.",
       bold_lead="R_i cross-activation-type scale caveat. ")
bullet("Report 1 (full capacity) never trained Hybrid Early or Hybrid Mid, so Part 3/4/5/8's "
       "full-capacity sections say nothing about them; only the reduced-capacity reports (which "
       "trained Hybrid Early, and have no Hybrid Mid analogue at depth 5) cover it.",
       bold_lead="Untrained Hybrid Early/Mid at full capacity. ")
bullet("Experiment 1 (a parameter/depth-matched non-equivariant SNN baseline) was never run at any "
       "capacity -- it requires building and training a new architecture, out of scope for every pass "
       "in this project so far, all of which were either representation-analysis-only or explicitly "
       "no-retraining.", bold_lead="Experiment 1 not run. ")
bullet("Report 3's 32-image fixed probe set quantizes accuracy-derived quantities (D_s in particular) "
       "in steps of 1/32 ~= 3.1%; this is why Experiment 2's D_s should be read as indicative, not "
       "precise, at this probe-set size.", bold_lead="Probe-set quantization. ")
bullet("The Fix 1 correction addresses the reference-norm-collapse artifact, but Experiment 3's "
       "original design goal (cleanly comparing D_{ℓ,t} at theta=45 vs. theta=90 to show the "
       "exact-group-angle guarantee reducing sensitivity specifically at 90 degrees) is still not "
       "achievable at the deepest layer even post-fix, since both angles are compared against the same "
       "reference representation and are equally exposed to any residual near-zero-norm cells.",
       bold_lead="D_{ℓ,t} 45-vs-90 comparison still limited at layer 5. ")
bullet("The claim that sparsity from an upstream spiking layer propagates into a downstream analog "
       "layer's collapse frequency (Part 6.3) is a plausible, mechanistically motivated reading of the "
       "corrected data, not a directly verified mechanism (e.g. via direct inspection of layer 4's "
       "pre-layer-5 output sparsity per configuration).",
       bold_lead="Upstream-sparsity-propagation mechanism unconfirmed. ")
bullet("Hybrid Early's large seed-2 accuracy instability (Part 2.2-2.3) was noted in both "
       "reduced-capacity regimes but never investigated further -- unclear if it is a property of that "
       "configuration at this capacity or an isolated bad initialization.",
       bold_lead="Hybrid Early seed-2 instability uninvestigated. ")

h2("11.2 Recommended Next Steps, Prioritized Against the Sept 19 Abstract / Sept 24 Full-Paper Deadlines")
p("Today's date: August 31, 2026 -- roughly 19 days to the abstract deadline, 24 to the full paper.",
  italic=True)

numbered("Both corrections (Fix 1, Fix 2) are already applied and this document already reflects "
         "them -- no further action needed before the abstract; simply draw the abstract's claims from "
         "Parts 5.3, 6, and 7 of this document rather than the superseded single-draw/fixed-epsilon "
         "numbers in the original Report 2/3 files.",
         bold_lead="Immediate (before Sept 19), zero additional compute: ")
numbered("Directly test the upstream-sparsity-propagation mechanism (Part 6.3/11.1) by inspecting layer "
         "4's pre-layer-5 output sparsity per configuration on existing checkpoints -- cheap, no "
         "retraining, would upgrade a plausible reading to a confirmed one before the full paper.",
         bold_lead="Cheap, before Sept 24: ")
numbered("If GPU time allows before the full-paper deadline, add 2 more model seeds (to n=5) for "
         "Experiment 2's E90 specifically, since Part 5.3 identified overlapping mean+/-std bands "
         "between adjacent configurations as the main remaining weakness in an otherwise "
         "cross-regime-consistent finding -- this is the single highest-leverage additional-compute "
         "item in the project given it directly strengthens the report's most novel corrected finding.",
         bold_lead="Moderate effort, before Sept 24 if time allows: ")
numbered("Complete reduced-capacity Regime B seed 2's representation-level analysis (a few-minutes-per-"
         "stage job) to bring Part 4/6/8's Regime B tables from n=2 to n=3, matching Regime A.",
         bold_lead="Moderate effort: ")
numbered("Train Hybrid Early and Hybrid Mid at full capacity, extending Parts 3-5/8's full-capacity "
         "sections to the same 5-configuration coverage the reduced-capacity reports already have; "
         "run a controlled ablation isolating which of P4BatchNorm2d/longer training/higher T explains "
         "the original compute-constrained draft's non-replicating findings (Part 2.1).",
         bold_lead="Larger effort, only if time remains after the above: ")
numbered("Experiment 1 (parameter/depth-matched non-equivariant SNN baseline) -- the largest remaining "
         "lift in the project, requiring a new architecture and full retraining; realistically a "
         "full-paper-cycle item rather than an abstract-cycle one.",
         bold_lead="Largest effort, full-paper timeframe only: ")

doc.save(os.path.join(BASE, "SNN_Rotation_Invariance_Consolidated_Report.docx"))
print("Saved SNN_Rotation_Invariance_Consolidated_Report.docx")



